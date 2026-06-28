"""
EcoDoc - Ekologik chiqindi hujjatlarini avtomatlashtiruvchi bot
report.py - Word (.docx) va PDF hisobotlar yaratish

Bu modul quyidagilarni qiladi:
1. AI yaratgan matnni professional Word hujjatga aylantiradi
2. Korxona ma'lumotlari, shtamp, imzo joylari qo'shadi
3. PDF variantini ham yaratadi
4. O'zbek va nemis dizaynlarini qo'llab-quvvatlaydi

Ausbildung uchun: Professional hujjat dizayni muhim!
"""

import os
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging

from config import UZ_CONFIG, DE_CONFIG, REPORTS_DIR

logger = logging.getLogger(__name__)


def create_word_report(
    language: str,
    company_name: str,
    company_address: str,
    company_inn: str,
    company_phone: str,
    waste_type: str,
    quantity: float,
    report_date: str,
    ai_generated_text: str,
    director_name: str = None,
    waste_category: str = None,
    report_id: int = None
) -> Optional[str]:
    """
    Professional Word hisobot hujjatini yaratish.

    Args:
        language: Til kodi ('uz' yoki 'de')
        company_name: Korxona nomi
        company_address: Korxona manzili
        company_inn: INN/Steuernummer
        company_phone: Telefon
        waste_type: Chiqindi turi
        quantity: Miqdori
        report_date: Sana
        ai_generated_text: AI yaratgan hisobot matni
        director_name: Direktor ismi
        waste_category: Chiqindi kategoriyasi (nemis uchun)
        report_id: Hisobot IDsi (fayl nomi uchun)

    Returns:
        str: Yaratilgan fayl yo'li yoki None
    """

    try:
        # Hisobotlar papkasini yaratish
        os.makedirs(REPORTS_DIR, exist_ok=True)

        # Tilga qarab konfiguratsiyani tanlash
        config = UZ_CONFIG if language == "uz" else DE_CONFIG

        # Word hujjatini yaratish
        doc = Document()

        # ==========================================
        # 1. SAHIFA SOZLAMALARI
        # ==========================================
        _setup_page(doc)

        # ==========================================
        # 2. SARIQ (HEADER)
        # ==========================================
        _add_header(doc, config, language)

        # ==========================================
        # 3. KORXONA MA'LUMOTLARI JADVALI
        # ==========================================
        _add_company_info_table(
            doc, config, company_name, company_address,
            company_inn, company_phone, director_name, report_date
        )

        # ==========================================
        # 4. CHIQINDI MA'LUMOTLARI
        # ==========================================
        _add_waste_info(
            doc, config, waste_type, quantity, waste_category
        )

        # ==========================================
        # 5. AI YARATGAN HISOBOT MATNI
        # ==========================================
        _add_ai_report_text(doc, ai_generated_text)

        # ==========================================
        # 6. IZOHLAR VA IMZO BO'LIMI
        # ==========================================
        _add_signature_section(doc, config, director_name)

        # ==========================================
        # 7. FAYLNI SAQLASH
        # ==========================================
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        lang_prefix = "UZ" if language == "uz" else "DE"
        filename = f"EcoDoc_{lang_prefix}_{report_id or timestamp}_{company_name[:20]}.docx"
        filepath = os.path.join(REPORTS_DIR, filename)

        doc.save(filepath)
        logger.info(f"✅ Word hisobot yaratildi: {filepath}")

        return filepath

    except Exception as e:
        logger.error(f"❌ Word hisobot yaratishda xatolik: {e}")
        return None


def _setup_page(doc: Document):
    """
    Word sahifasini sozlash (chegaralar, o'lcham).

    Args:
        doc: Word hujjati obyekti
    """
    # Har bir seksiya uchun chegaralar
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_header(doc: Document, config: dict, language: str):
    """
    Hujjat sarlavhasini qo'shish.

    Args:
        doc: Word hujjati
        config: Til konfiguratsiyasi
        language: Til kodi
    """
    # Asosiy sarlavha
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config["doc_title"])
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 100, 0)  # Yashil rang (ekologiya)

    # Qonunlar ro'yxati (kichik shrift)
    if language == "uz":
        laws_text = "Qonuniy asos: " + ", ".join(config["laws"][:2])
    else:
        laws_text = "Rechtliche Grundlage: " + ", ".join(config["laws"][:2])

    laws_para = doc.add_paragraph()
    laws_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    laws_run = laws_para.add_run(laws_text)
    laws_run.font.size = Pt(9)
    laws_run.font.color.rgb = RGBColor(100, 100, 100)
    laws_run.italic = True

    # Chiziq chizish
    doc.add_paragraph("_" * 50)


def _add_company_info_table(
    doc: Document,
    config: dict,
    company_name: str,
    company_address: str,
    company_inn: str,
    company_phone: str,
    director_name: str,
    report_date: str
):
    """
    Korxona ma'lumotlari jadvalini qo'shish.

    Args:
        doc: Word hujjati
        config: Til konfiguratsiyasi
        company_name: Korxona nomi
        company_address: Manzil
        company_inn: INN
        company_phone: Telefon
        director_name: Direktor
        report_date: Sana
    """
    # Jadval yaratish (2 ustun, 5 qator)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Jadval kengligini sozlash
    for row in table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)

    # Ma'lumotlarni to'ldirish
    data = [
        (config["organization_label"], company_name),
        ("Manzil / Adresse" if config.get("country") == "O'zbekiston" else "Adresse", company_address),
        ("INN / Steuernummer", company_inn or "-"),
        ("Telefon", company_phone or "-"),
        (config["responsible_label"], director_name or "-"),
    ]

    for i, (label, value) in enumerate(data):
        # Birinchi ustun (sarlavha)
        cell_label = table.rows[i].cells[0]
        cell_label.text = label
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

        # Ikkinchi ustun (qiymat)
        cell_value = table.rows[i].cells[1]
        cell_value.text = value
        for paragraph in cell_value.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()  # Bo'sh qator


def _add_waste_info(
    doc: Document,
    config: dict,
    waste_type: str,
    quantity: float,
    waste_category: str = None
):
    """
    Chiqindi ma'lumotlarini qo'shish.

    Args:
        doc: Word hujjati
        config: Til konfiguratsiyasi
        waste_type: Chiqindi turi
        quantity: Miqdori
        waste_category: Kategoriya (nemis uchun)
    """
    # Bo'lim sarlavhasi
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run("📋 " + ("CHIQINDI MA'LUMOTLARI" if config.get("country") == "O'zbekiston" else "ABFALLDATEN"))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 100, 0)

    # Jadval yaratish
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    data = [
        (config["waste_type_label"], waste_type),
        (config["quantity_label"], f"{quantity} kg"),
    ]

    if waste_category:
        data.append(("Kategorie / Kategoriya", waste_category))
    else:
        data.append(("Sana / Datum", datetime.now().strftime(config["date_format"])))

    for i, (label, value) in enumerate(data):
        cell_label = table.rows[i].cells[0]
        cell_label.text = label
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True

        cell_value = table.rows[i].cells[1]
        cell_value.text = value

    doc.add_paragraph()


def _add_ai_report_text(doc: Document, ai_text: str):
    """
    AI yaratgan hisobot matnini qo'shish.

    Args:
        doc: Word hujjati
        ai_text: AI matni
    """
    # Bo'lim sarlavhasi
    heading = doc.add_paragraph()
    run = heading.add_run("📝 HISOBOT MATNI / BERICHT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 100, 0)

    # AI matnini qo'shish
    paragraphs = ai_text.split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(para_text.strip())
            run.font.size = Pt(11)

    doc.add_paragraph()


def _add_signature_section(doc: Document, config: dict, director_name: str = None):
    """
    Imzo va sana bo'limini qo'shish.

    Args:
        doc: Word hujjati
        config: Til konfiguratsiyasi
        director_name: Direktor ismi
    """
    # Chiziq
    doc.add_paragraph("_" * 50)

    # Imzo jadvali (2 ustun)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Chap tomon - Imzo
    table.rows[0].cells[0].text = config["signature_label"]
    table.rows[1].cells[0].text = "_" * 20

    # O'ng tomon - Sana
    table.rows[0].cells[1].text = config["date_label"]
    table.rows[1].cells[1].text = datetime.now().strftime(config["date_format"])

    # Qalin shrift
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Mas'ul shaxs
    if director_name:
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"{config['responsible_label']}: {director_name}")
        run.font.size = Pt(10)
        run.italic = True


def add_watermark(doc: Document, text: str = "EcoDoc"):
    """
    Hujjatga suv belgisi qo'shish (professional ko'rinish uchun).

    Args:
        doc: Word hujjati
        text: Suv belgisi matni
    """
    # Bu funksiya keyinroq qo'shilishi mumkin
    # Word hujjatlarida suv belgisi murakkabroq
    pass


# Test uchun
if __name__ == "__main__":
    print("🧪 Hisobot yaratish testi...")

    test_text = """KIRISH
Bu hisobot Test Korxona tomonidan 2026-yil 15-iyun holatiga ko'ra chiqindi holatini yoritib beradi.

KORXONA HAQIDA
Test Korxona Toshkent shahrida joylashgan va plastik mahsulotlar ishlab chiqaradi.

CHIQINDI TAVSIFI
Plastik idishlar - ishlab chiqarish jarayonida hosil bo'lgan qoldiqlar."""

    result = create_word_report(
        language="uz",
        company_name="Test Korxona",
        company_address="Toshkent, O'zbekiston",
        company_inn="123456789",
        company_phone="+998 90 123 45 67",
        waste_type="Plastik idishlar",
        quantity=150.5,
        report_date="15.06.2026",
        ai_generated_text=test_text,
        director_name="Test Testov",
        report_id=1
    )

    if result:
        print(f"✅ Test hisobot yaratildi: {result}")
    else:
        print("❌ Test hisobot yaratilmadi")
